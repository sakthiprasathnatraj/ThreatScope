"""
ThreatScope Project Report - Word Document Generator
=====================================================
Run this script to generate a single Word (.docx) file from all chapter files.

Requirements:
    pip install python-docx

Usage:
    python generate_report.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Configuration ────────────────────────────────────────────────────────────
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(REPORT_DIR, "ThreatScope_Project_Report.docx")

CHAPTER_FILES = [
    "CHAPTER-01-ABSTRACT-INTRODUCTION.md",
    "CHAPTER-02-SYSTEM-ANALYSIS.md",
    "CHAPTER-03-SYSTEM-REQUIREMENTS.md",
    "CHAPTER-04-SYSTEM-OVERVIEW.md",
    "CHAPTER-05-SYSTEM-DESIGN.md",
    "CHAPTER-06-TESTING-IMPLEMENTATION.md",
    "CHAPTER-07-08-CONCLUSION-FUTURE.md",
    "CHAPTER-09-10-REFERENCES-APPENDIX.md",
]


# ─── Helper: Set paragraph spacing ───────────────────────────────────────────
def set_spacing(paragraph, before=6, after=6, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


# ─── Helper: Add horizontal rule ─────────────────────────────────────────────
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_spacing(p, 0, 0)


# ─── Helper: Style a run ──────────────────────────────────────────────────────
def style_run(run, bold=False, italic=False, color=None, size=None, font_name="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)


# ─── Title Page ───────────────────────────────────────────────────────────────
def add_title_page(doc):
    # College name placeholder
    college = doc.add_paragraph()
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = college.add_run("[YOUR COLLEGE NAME]")
    style_run(r, bold=True, size=16, color=(31, 73, 125))
    set_spacing(college, before=72, after=6)

    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dept.add_run("Department of Computer Science and Engineering")
    style_run(r, size=13, color=(68, 114, 196))
    set_spacing(dept, before=4, after=30)

    add_horizontal_rule(doc)

    # Project title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("THREATSCOPE")
    style_run(r, bold=True, size=28, color=(31, 73, 125))
    set_spacing(title, before=36, after=6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("A User-Centric Real-Time Network Threat Detection\nand Explanation Platform")
    style_run(r, bold=True, size=14, color=(68, 114, 196))
    set_spacing(subtitle, before=4, after=30)

    add_horizontal_rule(doc)

    # Submission info
    info_lines = [
        ("A Project Report submitted in partial fulfilment of the requirements", False),
        ("for the award of the degree of", False),
        ("", False),
        ("Bachelor of Engineering / Bachelor of Technology", True),
        ("in Computer Science and Engineering", True),
        ("", False),
        ("Submitted By", False),
        ("[Your Name]  |  Reg. No: [Your Register Number]", True),
        ("", False),
        ("Under the Guidance of", False),
        ("[Guide Name], [Designation]", True),
        ("", False),
        ("Academic Year: 2025–2026", False),
    ]

    for line, bold in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        style_run(r, bold=bold, size=12 if bold else 11)
        set_spacing(p, before=3, after=3)

    doc.add_page_break()


# ─── Table of Contents (manual) ───────────────────────────────────────────────
def add_table_of_contents(doc):
    h = doc.add_heading("TABLE OF CONTENTS", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(h, before=0, after=12)

    toc_entries = [
        ("ABSTRACT", ""),
        ("CHAPTER 1 – INTRODUCTION", ""),
        ("    1.1  Background and Motivation", ""),
        ("    1.2  Problem Statement", ""),
        ("    1.3  Scope of the Project", ""),
        ("    1.4  Organisation of the Report", ""),
        ("CHAPTER 2 – SYSTEM ANALYSIS", ""),
        ("    2.1  Study of Existing Systems", ""),
        ("    2.2  Problem Analysis and Proposed Solution", ""),
        ("CHAPTER 3 – SYSTEM REQUIREMENTS", ""),
        ("    3.1  Requirements Specification", ""),
        ("CHAPTER 4 – SYSTEM OVERVIEW", ""),
        ("    4.1  System Architecture Overview", ""),
        ("CHAPTER 5 – SYSTEM DESIGN", ""),
        ("    5.1  Detailed System Design", ""),
        ("CHAPTER 6 – SYSTEM TESTING AND IMPLEMENTATION", ""),
        ("    6.1  Testing Strategy and Results", ""),
        ("CHAPTER 7 – CONCLUSION", ""),
        ("    7.1  Summary of Achievements", ""),
        ("CHAPTER 8 – FUTURE ENHANCEMENT", ""),
        ("    8.1  Planned Improvements", ""),
        ("CHAPTER 9 – REFERENCES", ""),
        ("    9.1  References", ""),
        ("CHAPTER 10 – APPENDIX", ""),
        ("    10.1 Source Code Listings", ""),
    ]

    for entry, _ in toc_entries:
        p = doc.add_paragraph(entry)
        is_chapter = entry.startswith("CHAPTER") or entry == "ABSTRACT"
        run = p.runs[0] if p.runs else p.add_run(entry)
        run.bold = is_chapter
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        set_spacing(p, before=2, after=2)

    doc.add_page_break()


# ─── Markdown Parser & Renderer ───────────────────────────────────────────────
def parse_and_add_markdown(doc, md_text):
    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []
    in_table = False
    table_obj = None
    table_headers = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Code block ──────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Add code block as styled paragraph
                for cl in code_lines:
                    p = doc.add_paragraph(cl if cl else " ")
                    p.style = "No Spacing"
                    run = p.runs[0] if p.runs else p.add_run(cl)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                    # Light grey background via shading
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F0F0F0')
                    pPr.append(shd)
                    set_spacing(p, 0, 0)
                doc.add_paragraph()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────
        if line.strip() in ("---", "===", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Heading levels ───────────────────────────────────
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p, before=18, after=8)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            set_spacing(p, before=14, after=6)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            set_spacing(p, before=10, after=4)
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            set_spacing(p, before=8, after=4)
            i += 1
            continue

        # ── Table detection ──────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            # Collect all consecutive table lines
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Filter out separator rows (---|---)
            data_rows = [r for r in table_lines if not re.match(r'^\|[\s\-\|:]+\|$', r.strip())]
            if not data_rows:
                continue

            # Parse cells
            parsed_rows = []
            for row in data_rows:
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                parsed_rows.append(cells)

            if not parsed_rows:
                continue

            cols = max(len(r) for r in parsed_rows)
            # Pad rows
            for r in parsed_rows:
                while len(r) < cols:
                    r.append("")

            tbl = doc.add_table(rows=len(parsed_rows), cols=cols)
            tbl.style = "Table Grid"

            for ri, row in enumerate(parsed_rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
                        # Header row background
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), '4472C4')
                        tcPr.append(shd)
                        run.font.color.rgb = RGBColor(255, 255, 255)

            doc.add_paragraph()
            continue

        # ── Bullet / list items ──────────────────────────────
        if re.match(r'^[\-\*] ', line):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            set_spacing(p, before=2, after=2)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            set_spacing(p, before=2, after=2)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(89, 89, 89)
            set_spacing(p, before=4, after=4)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph with inline formatting ───────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Process inline **bold** and `code`
        segments = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
        for seg in segments:
            if seg.startswith("**") and seg.endswith("**"):
                run = p.add_run(seg[2:-2])
                style_run(run, bold=True, size=11)
            elif seg.startswith("`") and seg.endswith("`"):
                run = p.add_run(seg[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 128)
            else:
                run = p.add_run(seg)
                style_run(run, size=11)

        set_spacing(p, before=3, after=3, line=14)
        i += 1


# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  ThreatScope - Word Report Generator")
    print("=" * 60)

    # Install python-docx if needed
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        from docx import Document

    doc = Document()

    # ── Page setup (A4) ───────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title Page ────────────────────────────────────────────
    print("→ Adding title page...")
    add_title_page(doc)

    # ── Table of Contents ─────────────────────────────────────
    print("→ Adding table of contents...")
    add_table_of_contents(doc)

    # ── Chapters ──────────────────────────────────────────────
    for fname in CHAPTER_FILES:
        fpath = os.path.join(REPORT_DIR, fname)
        if not os.path.exists(fpath):
            print(f"  ⚠ File not found: {fname}, skipping.")
            continue

        print(f"→ Processing {fname}...")
        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()

        parse_and_add_markdown(doc, content)
        doc.add_page_break()

    # ── Save ──────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print()
    print("=" * 60)
    print(f"  ✅ Report saved to:")
    print(f"  {OUTPUT_FILE}")
    print("=" * 60)


if __name__ == "__main__":
    main()
